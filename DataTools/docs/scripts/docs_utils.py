"""Shared utilities for DataTools documentation generators."""
from __future__ import annotations

import base64
import html as _html
import io
import re
from pathlib import Path


def embed_images_in_markdown(md: str, docs_root: Path) -> str:
    """Replace relative image links with base64 data URIs.

    Scans for Markdown image links of the form ``![alt](../screenshots/...)``
    and replaces them with inline base64 data URIs so the document is
    self-contained and renders correctly when pasted into Confluence or other
    external tools that cannot resolve relative file paths.

    Args:
        md:        Markdown text to process.
        docs_root: Absolute path to the ``docs/`` directory (parent of
                   ``generated/`` and ``screenshots/``).

    Returns:
        Markdown text with all resolvable PNG links replaced by data URIs.
    """
    generated_dir = docs_root / "generated"

    def _flatten(png_data: bytes) -> bytes:
        """Composite transparent PNG over a dark background (matches the app theme)."""
        try:
            from PIL import Image
            import io as _io
            img = Image.open(_io.BytesIO(png_data)).convert("RGBA")
            bg = Image.new("RGBA", img.size, (26, 26, 26, 255))
            bg.paste(img, mask=img.split()[3])
            buf = _io.BytesIO()
            bg.convert("RGB").save(buf, format="PNG")
            return buf.getvalue()
        except Exception:
            return png_data  # fall back if PIL unavailable

    def _replace(m: re.Match) -> str:
        alt = m.group(1)
        rel = m.group(2)
        abs_path = (generated_dir / rel).resolve()
        if abs_path.exists() and abs_path.suffix.lower() == ".png":
            data = base64.b64encode(_flatten(abs_path.read_bytes())).decode()
            return f"![{alt}](data:image/png;base64,{data})"
        return m.group(0)  # leave unchanged if file not found

    return re.sub(r"!\[([^\]]*)\]\(([^)]+\.png)\)", _replace, md)


# ---------------------------------------------------------------------------
# Markdown → HTML conversion
# ---------------------------------------------------------------------------

_HTML_STYLE = """
body { font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif;
       max-width: 960px; margin: 40px auto; padding: 0 24px;
       color: #172b4d; line-height: 1.6; }
h1 { font-size: 2em; border-bottom: 1px solid #dfe1e6; padding-bottom: .3em; }
h2 { font-size: 1.5em; border-bottom: 1px solid #dfe1e6; padding-bottom: .2em; margin-top: 1.8em; }
h3 { font-size: 1.15em; margin-top: 1.4em; }
h4 { font-size: 1em; margin-top: 1.2em; }
img { max-width: 100%; border: 1px solid #dfe1e6; border-radius: 4px; margin: 12px 0; }
table { border-collapse: collapse; width: 100%; margin: 16px 0; }
th, td { border: 1px solid #dfe1e6; padding: 8px 12px; text-align: left; }
th { background: #f4f5f7; font-weight: 600; }
tr:nth-child(even) td { background: #fafbfc; }
code { background: #f4f5f7; padding: 2px 5px; border-radius: 3px; font-size: .9em; }
pre { background: #f4f5f7; padding: 16px; border-radius: 4px; overflow-x: auto; }
pre code { background: none; padding: 0; }
blockquote { border-left: 4px solid #dfe1e6; margin: 0; padding: 0 16px; color: #5e6c84; }
ul, ol { padding-left: 1.5em; }
li { margin: .25em 0; }
a { color: #0052cc; }
"""


def markdown_to_html(md: str, title: str = "DataTools") -> str:
    """Convert a DataTools-generated Markdown document to a self-contained HTML file.

    Handles the subset of Markdown produced by the documentation generators:
    headings, paragraphs, bold/italic/inline-code, fenced code blocks,
    bullet/numbered lists, tables, images (including base64 data URIs), and links.

    Args:
        md:    Markdown source text (may contain base64 image data URIs).
        title: Used in the HTML ``<title>`` tag.

    Returns:
        Complete HTML document as a string.
    """
    lines = md.splitlines()
    out: list[str] = []
    i = 0
    in_list: str | None = None   # "ul" or "ol"
    in_table = False
    table_header_done = False

    def _flush_list():
        nonlocal in_list
        if in_list:
            out.append(f"</{in_list}>")
            in_list = None

    def _flush_table():
        nonlocal in_table, table_header_done
        if in_table:
            out.append("</tbody></table>")
            in_table = False
            table_header_done = False

    def _inline(text: str) -> str:
        """Apply inline markdown: bold, italic, inline-code, links, images."""
        # Images first (may contain base64)
        text = re.sub(
            r"!\[([^\]]*)\]\((data:image/[^)]+|[^)]+\.png)\)",
            lambda m: f'<img alt="{_html.escape(m.group(1))}" src="{m.group(2)}">',
            text,
        )
        # Links
        text = re.sub(
            r"\[([^\]]+)\]\(([^)]+)\)",
            lambda m: f'<a href="{m.group(2)}">{m.group(1)}</a>',
            text,
        )
        # Inline code (before bold/italic so backticks are safe)
        parts = re.split(r"`([^`]+)`", text)
        result = ""
        for j, part in enumerate(parts):
            if j % 2 == 1:
                result += f"<code>{_html.escape(part)}</code>"
            else:
                p = part
                p = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", p)
                p = re.sub(r"\*(.+?)\*", r"<em>\1</em>", p)
                result += p
        return result

    while i < len(lines):
        line = lines[i]

        # Fenced code block
        if line.startswith("```"):
            _flush_list()
            _flush_table()
            lang = line[3:].strip()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(_html.escape(lines[i]))
                i += 1
            out.append(f'<pre><code class="language-{lang}">' + "\n".join(code_lines) + "</code></pre>")
            i += 1
            continue

        # Headings
        m = re.match(r"(#{1,4})\s+(.*)", line)
        if m:
            _flush_list()
            _flush_table()
            level = len(m.group(1))
            out.append(f"<h{level}>{_inline(m.group(2))}</h{level}>")
            i += 1
            continue

        # Table row
        if line.startswith("|") and line.endswith("|"):
            cells = [c.strip() for c in line.strip("|").split("|")]
            # Separator row
            if all(re.match(r"^-+$", c.strip("-").strip()) or c.strip() == "" for c in cells):
                i += 1
                continue
            _flush_list()
            if not in_table:
                out.append('<table><thead><tr>')
                out.append("".join(f"<th>{_inline(c)}</th>" for c in cells))
                out.append("</tr></thead><tbody>")
                in_table = True
                table_header_done = True
            else:
                out.append("<tr>" + "".join(f"<td>{_inline(c)}</td>" for c in cells) + "</tr>")
            i += 1
            continue

        # Bullet list
        m = re.match(r"^[-*]\s+(.*)", line)
        if m:
            _flush_table()
            if in_list != "ul":
                _flush_list()
                out.append("<ul>")
                in_list = "ul"
            out.append(f"<li>{_inline(m.group(1))}</li>")
            i += 1
            continue

        # Numbered list
        m = re.match(r"^\d+\.\s+(.*)", line)
        if m:
            _flush_table()
            if in_list != "ol":
                _flush_list()
                out.append("<ol>")
                in_list = "ol"
            out.append(f"<li>{_inline(m.group(1))}</li>")
            i += 1
            continue

        # Blank line
        if line.strip() == "":
            _flush_list()
            _flush_table()
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^---+$", line.strip()):
            _flush_list()
            _flush_table()
            out.append("<hr>")
            i += 1
            continue

        # Plain paragraph line
        _flush_list()
        _flush_table()
        out.append(f"<p>{_inline(line)}</p>")
        i += 1

    _flush_list()
    _flush_table()

    body = "\n".join(out)
    confluence_note = (
        "<div style='background:#fffbe6;border:1px solid #f0c040;border-radius:4px;"
        "padding:10px 16px;margin-bottom:24px;font-size:.9em;color:#594300;'>"
        "<strong>Confluence import:</strong> In Confluence, use "
        "<em>Insert → Markup → HTML</em> and paste the full source of this file "
        "(Ctrl+U → Ctrl+A → Ctrl+C in browser, or open the .html file in a text editor). "
        "Do <strong>not</strong> copy from the browser page — images will not transfer."
        "</div>"
    )
    return (
        f"<!DOCTYPE html>\n<html lang='en'>\n<head>\n"
        f"<meta charset='utf-8'>\n<title>{_html.escape(title)}</title>\n"
        f"<style>{_HTML_STYLE}</style>\n</head>\n<body>\n{confluence_note}\n{body}\n</body>\n</html>\n"
    )


    return re.sub(r"!\[([^\]]*)\]\(([^)]+\.png)\)", _replace, md)


# ---------------------------------------------------------------------------
# Markdown → Word (.docx) conversion  — for Confluence import
# ---------------------------------------------------------------------------

def _flatten_png(png_data: bytes) -> bytes:
    """Composite transparent PNG over a dark background (matches app theme)."""
    try:
        from PIL import Image
        img = Image.open(io.BytesIO(png_data)).convert("RGBA")
        bg = Image.new("RGBA", img.size, (26, 26, 26, 255))
        bg.paste(img, mask=img.split()[3])
        buf = io.BytesIO()
        bg.convert("RGB").save(buf, format="PNG")
        return buf.getvalue()
    except Exception:
        return png_data


def markdown_to_docx(md: str, docs_root: Path, title: str = "DataTools") -> bytes:
    """Convert DataTools-generated Markdown to a Word .docx file (as bytes).

    Images are embedded directly from the screenshot folder so the document
    is self-contained. Word files can be imported into Confluence via
    *Space Tools → Content Tools → Import* and images will be preserved.

    Args:
        md:        Markdown source text with relative ``../screenshots/...`` links.
        docs_root: Absolute path to the ``docs/`` directory.
        title:     Document title written as the first heading if not already present.

    Returns:
        Raw bytes of the generated .docx file.
    """
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    generated_dir = docs_root / "generated"

    doc = Document()

    # ── Style helpers ────────────────────────────────────────────────────────
    def _add_heading(text: str, level: int):
        doc.add_heading(text, level=level)

    def _add_image(abs_path: Path):
        try:
            flat = _flatten_png(abs_path.read_bytes())
            doc.add_picture(io.BytesIO(flat), width=Inches(5.5))
        except Exception:
            doc.add_paragraph(f"[Image: {abs_path.name}]")

    def _apply_inline(para, text: str):
        """Add inline-formatted runs (bold, italic, code) to a paragraph."""
        # Strip image tags — handled separately
        text = re.sub(r"!\[[^\]]*\]\([^)]+\)", "", text)
        # Links: keep text only
        text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
        # Split on bold/italic/code markers
        token_re = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
        for part in token_re.split(text):
            if not part:
                continue
            run = para.add_run()
            if part.startswith("**") and part.endswith("**"):
                run.bold = True
                run.text = part[2:-2]
            elif part.startswith("*") and part.endswith("*"):
                run.italic = True
                run.text = part[1:-1]
            elif part.startswith("`") and part.endswith("`"):
                run.font.name = "Courier New"
                run.font.size = Pt(9)
                run.text = part[1:-1]
            else:
                run.text = part

    # ── Parse lines ──────────────────────────────────────────────────────────
    lines = md.splitlines()
    i = 0
    in_table = False
    tbl = None
    table_cols: list[str] = []

    def _flush_table():
        nonlocal in_table, tbl, table_cols
        in_table = False
        tbl = None
        table_cols = []

    while i < len(lines):
        line = lines[i]

        # Fenced code block
        if line.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            p = doc.add_paragraph("\n".join(code_lines), style="No Spacing")
            p.runs[0].font.name = "Courier New"
            p.runs[0].font.size = Pt(8)
            i += 1
            _flush_table()
            continue

        # Headings
        m = re.match(r"(#{1,4})\s+(.*)", line)
        if m:
            _flush_table()
            level = len(m.group(1))
            _add_heading(m.group(2), level)
            i += 1
            continue

        # Image line
        m = re.match(r"!\[([^\]]*)\]\(([^)]+\.png)\)", line.strip())
        if m:
            _flush_table()
            rel = m.group(2)
            abs_path = (generated_dir / rel).resolve()
            if abs_path.exists():
                _add_image(abs_path)
            i += 1
            continue

        # Table
        if line.startswith("|") and line.endswith("|"):
            cells = [c.strip() for c in line.strip("|").split("|")]
            # Separator row
            if all(re.match(r"^-+$", c.strip("-").strip()) or not c.strip() for c in cells):
                i += 1
                continue
            if not in_table:
                in_table = True
                table_cols = cells
                tbl = doc.add_table(rows=1, cols=len(cells))
                tbl.style = "Table Grid"
                hdr = tbl.rows[0].cells
                for j, col in enumerate(cells):
                    hdr[j].text = col
                    for run in hdr[j].paragraphs[0].runs:
                        run.bold = True
            else:
                row = tbl.add_row().cells
                for j, col in enumerate(cells[:len(table_cols)]):
                    row[j].text = col
            i += 1
            continue

        # Bullet list
        m = re.match(r"^[-*]\s+(.*)", line)
        if m:
            _flush_table()
            p = doc.add_paragraph(style="List Bullet")
            _apply_inline(p, m.group(1))
            i += 1
            continue

        # Numbered list
        m = re.match(r"^\d+\.\s+(.*)", line)
        if m:
            _flush_table()
            p = doc.add_paragraph(style="List Number")
            _apply_inline(p, m.group(1))
            i += 1
            continue

        # Blank line
        if not line.strip():
            _flush_table()
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^---+$", line.strip()):
            _flush_table()
            i += 1
            continue

        # Plain paragraph
        _flush_table()
        p = doc.add_paragraph()
        _apply_inline(p, line)
        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

